import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_styled_notes_docx(md_path, docx_path):
    doc = Document()

    # Page Margins: 1 inch
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Palette
    COLOR_TEAL = RGBColor(13, 79, 79)       # #0D4F4F
    COLOR_CYAN = RGBColor(0, 180, 216)      # #00B4D8
    COLOR_DARK = RGBColor(26, 26, 46)       # #1A1A2E
    COLOR_AMBER = RGBColor(247, 127, 0)     # #F77F00
    COLOR_GRAY = RGBColor(100, 100, 100)

    # Styles
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(51, 51, 51)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False

    for line in lines:
        raw_line = line.rstrip('\r\n')

        # Code block handling
        if raw_line.startswith('```'):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(raw_line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = COLOR_DARK
            continue

        # Horizontal rule
        if raw_line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            continue

        # Headings
        if raw_line.startswith('# '):
            h = doc.add_heading(level=1)
            h.paragraph_format.space_before = Pt(16)
            h.paragraph_format.space_after = Pt(6)
            r = h.add_run(raw_line[2:])
            r.font.name = 'Calibri'
            r.font.size = Pt(20)
            r.font.bold = True
            r.font.color.rgb = COLOR_TEAL
        elif raw_line.startswith('## '):
            h = doc.add_heading(level=2)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(4)
            r = h.add_run(raw_line[3:])
            r.font.name = 'Calibri'
            r.font.size = Pt(14)
            r.font.bold = True
            r.font.color.rgb = COLOR_TEAL
        elif raw_line.startswith('### '):
            h = doc.add_heading(level=3)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(2)
            r = h.add_run(raw_line[4:])
            r.font.name = 'Calibri'
            r.font.size = Pt(12)
            r.font.bold = True
            r.font.color.rgb = COLOR_DARK
        elif raw_line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            content = raw_line[2:]
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                    r.font.color.rgb = COLOR_DARK
                elif part.startswith('*') and part.endswith('*'):
                    r = p.add_run(part[1:-1])
                    r.italic = True
                else:
                    r = p.add_run(part)
                r.font.name = 'Calibri'
                r.font.size = Pt(11)
        elif raw_line.startswith('"') or raw_line.startswith('**['):
            # Speaking script callout box
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            cell = table.cell(0, 0)
            cell.width = Inches(6.5)
            
            shading_elm = parse_xml(r'<w:shd {} w:fill="F0F9FA"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            borders_elm = parse_xml(r'''
                <w:tcBorders {} >
                    <w:top w:val="none"/>
                    <w:left w:val="single" w:sz="24" w:space="0" w:color="00B4D8"/>
                    <w:bottom w:val="none"/>
                    <w:right w:val="none"/>
                </w:tcBorders>
            '''.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(borders_elm)

            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', raw_line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                    r.font.color.rgb = COLOR_TEAL
                elif part.startswith('*') and part.endswith('*'):
                    r = p.add_run(part[1:-1])
                    r.italic = True
                    r.font.color.rgb = COLOR_DARK
                else:
                    r = p.add_run(part)
                    r.font.color.rgb = COLOR_DARK
                r.font.name = 'Calibri'
                r.font.size = Pt(10.5)
            
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        elif raw_line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', raw_line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                    r.font.color.rgb = COLOR_TEAL
                elif part.startswith('*') and part.endswith('*'):
                    r = p.add_run(part[1:-1])
                    r.italic = True
                else:
                    r = p.add_run(part)
                r.font.name = 'Calibri'
                r.font.size = Pt(11)

    doc.save(docx_path)
    print(f"Successfully generated DOCX: {docx_path}")

if __name__ == '__main__':
    md = r'D:\ICBT\Final\digital-therapy-agent\presentation\FINAL_VIVA_SPEAKER_NOTES.md'
    docx = r'D:\ICBT\Final\digital-therapy-agent\presentation\FINAL_VIVA_SPEAKER_NOTES.docx'
    create_styled_notes_docx(md, docx)
